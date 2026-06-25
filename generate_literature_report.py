"""
生成文献学习综合分析报告（Word格式）
"""

import os
import sys
import json
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    sys.exit(1)

# ============================================================
# 配置
# ============================================================
DATA_DIR = r"D:\VScode\firstcc\GHGs-WWTPs\output\literature_learning"
KNOWLEDGE_DIR = r"D:\VScode\firstcc\GHGs-WWTPs\knowledge_store"
OUTPUT_DIR = r"D:\VScode\firstcc\GHGs-WWTPs\output\literature_learning"

# ============================================================
# 辅助函数
# ============================================================

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E75B6')

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            # 交替行颜色
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'D6E4F0')

    return table


def add_heading_with_number(doc, text, level):
    """添加带编号的标题"""
    heading = doc.add_heading(text, level=level)
    return heading


# ============================================================
# 主报告生成
# ============================================================

def generate_report():
    """生成完整的Word报告"""

    # 加载数据
    with open(os.path.join(DATA_DIR, "all_papers_analysis.json"), "r", encoding="utf-8") as f:
        all_results = json.load(f)

    with open(os.path.join(KNOWLEDGE_DIR, "learned_writing_patterns.json"), "r", encoding="utf-8") as f:
        writing_patterns = json.load(f)

    with open(os.path.join(KNOWLEDGE_DIR, "learned_analysis_methods.json"), "r", encoding="utf-8") as f:
        analysis_methods = json.load(f)

    with open(os.path.join(KNOWLEDGE_DIR, "learned_figure_design.json"), "r", encoding="utf-8") as f:
        figure_design = json.load(f)

    # 创建文档
    doc = Document()

    # ============================================================
    # 封面
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("污水处理厂温室气体排放文献\n综合分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("——写作模式、数据分析方法与图表设计学习总结")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # 信息框
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"分析文献数量：{len(all_results)} 篇\n")
    run.font.size = Pt(12)
    run = info_para.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}\n")
    run.font.size = Pt(12)
    run = info_para.add_run("分析工具：GHGs-WWTPs 文献学习系统")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ============================================================
    # 目录页
    # ============================================================
    doc.add_heading("目  录", level=1)
    doc.add_paragraph()

    toc_items = [
        ("一、研究概述", 3),
        ("    1.1 研究背景与目的", 3),
        ("    1.2 文献来源与范围", 3),
        ("    1.3 分析方法", 3),
        ("二、写作模式分析", 4),
        ("    2.1 高频过渡词分析", 4),
        ("    2.2 学术模糊语（Hedging）使用", 5),
        ("    2.3 强调语使用分析", 5),
        ("    2.4 数据报告句式", 6),
        ("    2.5 引用格式统计", 6),
        ("三、数据分析方法汇总", 7),
        ("    3.1 统计检验方法", 7),
        ("    3.2 回归分析方法", 8),
        ("    3.3 机器学习方法", 8),
        ("    3.4 不确定性分析方法", 9),
        ("    3.5 排放核算方法", 9),
        ("    3.6 数据处理方法", 10),
        ("四、图表使用分析", 11),
        ("    4.1 常用图表类型", 11),
        ("    4.2 图表数量统计", 11),
        ("    4.3 表格复杂度分析", 12),
        ("五、写作经验总结", 13),
        ("    5.1 摘要写作要点", 13),
        ("    5.2 引言写作模式", 13),
        ("    5.3 方法描述规范", 14),
        ("    5.4 结果呈现技巧", 14),
        ("    5.5 讨论写作框架", 15),
        ("六、文献列表", 16),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item} {'.' * (50 - len(item))} {page}")
        run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # 第一章：研究概述
    # ============================================================
    doc.add_heading("一、研究概述", level=1)

    doc.add_heading("1.1 研究背景与目的", level=2)
    doc.add_paragraph(
        "污水处理厂（WWTPs）是温室气体（GHG）排放的重要来源之一，主要包括二氧化碳（CO2）、"
        "甲烷（CH4）和氧化亚氮（N2O）。准确核算和有效控制污水处理过程中的温室气体排放，"
        "对于实现碳中和目标具有重要意义。"
    )
    doc.add_paragraph(
        "本报告通过对117篇相关文献进行系统分析，旨在："
    )
    doc.add_paragraph("（1）总结文献中常用的写作模式和学术表达方式；", style='List Bullet')
    doc.add_paragraph("（2）梳理文献中使用的数据分析方法及其应用频率；", style='List Bullet')
    doc.add_paragraph("（3）归纳图表设计和数据可视化的方法与技巧；", style='List Bullet')
    doc.add_paragraph("（4）为后续论文写作提供经验参考和方法指导。", style='List Bullet')

    doc.add_heading("1.2 文献来源与范围", level=2)
    doc.add_paragraph(
        f"本次分析共收录 {len(all_results)} 篇文献，主要来源于Web of Science等学术数据库，"
        "涵盖2020-2026年间发表的高水平期刊论文。文献主题涵盖污水处理厂温室气体排放的核算方法、"
        "影响因素、减排技术等多个方面。"
    )

    # 文献年份分布
    year_counts = {}
    for r in all_results:
        year = r.get("metadata", {}).get("year", "未知")
        if year and year != "未知":
            try:
                year = int(year)
                if 2000 <= year <= 2030:
                    year_counts[year] = year_counts.get(year, 0) + 1
            except:
                pass

    doc.add_paragraph("文献年份分布如下表所示：")
    year_rows = sorted(year_counts.items(), key=lambda x: x[0])
    add_styled_table(doc, ["年份", "文献数量", "占比"],
                     [[y, c, f"{c/len(all_results)*100:.1f}%"] for y, c in year_rows])

    doc.add_heading("1.3 分析方法", level=2)
    doc.add_paragraph(
        "本报告采用以下分析方法："
    )
    doc.add_paragraph("文本挖掘：提取文献中的关键词、过渡词、引用格式等信息；", style='List Number')
    doc.add_paragraph("模式识别：识别写作模式、数据报告方式、图表类型等；", style='List Number')
    doc.add_paragraph("统计分析：对各类方法的使用频率进行统计和排序；", style='List Number')
    doc.add_paragraph("经验总结：基于统计数据，提炼写作经验和最佳实践。", style='List Number')

    doc.add_page_break()

    # ============================================================
    # 第二章：写作模式分析
    # ============================================================
    doc.add_heading("二、写作模式分析", level=1)

    doc.add_heading("2.1 高频过渡词分析", level=2)
    doc.add_paragraph(
        "过渡词是学术写作中连接句子和段落的重要工具。通过对117篇文献的分析，"
        "我们统计了各类过渡词的使用频率，为后续写作提供参考。"
    )

    # 过渡词表格
    transition_words = writing_patterns.get("transition_words", {})
    tw_sorted = sorted(transition_words.items(), key=lambda x: -x[1])[:20]

    doc.add_paragraph("表2-1 高频过渡词统计（Top 20）")
    add_styled_table(doc, ["排名", "过渡词", "出现总次数", "平均频率（次/篇）"],
                     [[i+1, w, c, f"{c/len(all_results):.1f}"] for i, (w, c) in enumerate(tw_sorted)])

    doc.add_paragraph()
    doc.add_paragraph(
        "分析结果显示，\"however\"是使用频率最高的过渡词（平均5.7次/篇），"
        "表明学术写作中对比和转折是常见的逻辑关系。\"therefore\"和\"thus\"的高频率"
        "说明因果关系的表达在论文中同样重要。"
    )

    doc.add_heading("2.2 学术模糊语（Hedging）使用", level=2)
    doc.add_paragraph(
        "学术模糊语是学术写作中的重要特征，用于表达谨慎的态度和避免过于绝对的结论。"
        "常见的模糊语包括may、could、suggest、indicate等。"
    )

    hedging_phrases = writing_patterns.get("hedging_phrases", {})
    hp_sorted = sorted(hedging_phrases.items(), key=lambda x: -x[1])[:15]

    doc.add_paragraph("表2-2 学术模糊语使用频率")
    add_styled_table(doc, ["模糊语", "出现次数", "平均频率（次/篇）"],
                     [[w, c, f"{c/len(all_results):.1f}"] for w, c in hp_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "\"could\"和\"may\"是使用最频繁的模糊语，分别达到6.4次/篇和5.7次/篇。"
        "这反映了学术写作中对不确定性的谨慎表达。\"estimated\"和\"approximately\""
        "等词的高频使用也表明，数据估算和近似表达是该领域的常见做法。"
    )

    doc.add_heading("2.3 强调语使用分析", level=2)

    emphasis_phrases = writing_patterns.get("emphasis_phrases", {})
    ep_sorted = sorted(emphasis_phrases.items(), key=lambda x: -x[1])[:10]

    doc.add_paragraph("表2-3 强调语使用频率")
    add_styled_table(doc, ["强调语", "出现次数", "平均频率（次/篇）"],
                     [[w, c, f"{c/len(all_results):.1f}"] for w, c in ep_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "\"significantly\"是使用最多的强调语（2.9次/篇），这与统计检验结果的报告密切相关。"
        "\"especially\"和\"particularly\"常用于突出特定发现或条件。"
    )

    doc.add_heading("2.4 数据报告句式", level=2)
    doc.add_paragraph(
        "通过对文献中数据报告方式的分析，我们发现以下常见模式："
    )

    data_patterns = [
        ("均值±标准差", "183.88 ± 2525.18", "描述数据集中趋势和离散程度"),
        ("百分比", "93%、68%、49%", "报告比例或贡献率"),
        ("相关系数", "r = 0.621、r = 0.55", "报告变量间相关性"),
        ("显著性水平", "P<0.05、p <0.05", "报告统计检验结果"),
        ("温度/浓度", "1°C、30.5%", "报告环境参数"),
    ]

    doc.add_paragraph("表2-4 常见数据报告句式")
    add_styled_table(doc, ["报告类型", "示例", "用途说明"],
                     [[t, e, u] for t, e, u in data_patterns])

    doc.add_heading("2.5 引用格式统计", level=2)

    citation_patterns = writing_patterns.get("citation_patterns", [])
    citation_summary = {}
    for cp in citation_patterns:
        pattern = cp.get("pattern", "")
        count = cp.get("count", 0)
        if "A-Z" in pattern or "Author" in pattern:
            citation_summary["Author (Year)"] = citation_summary.get("Author (Year)", 0) + count
        elif "\\d" in pattern:
            citation_summary["[Number]"] = citation_summary.get("[Number]", 0) + count

    doc.add_paragraph("表2-5 引用格式统计")
    citation_rows = [[fmt, count, f"{count/sum(citation_summary.values())*100:.1f}%"]
                     for fmt, count in sorted(citation_summary.items(), key=lambda x: -x[1])]
    add_styled_table(doc, ["引用格式", "使用次数", "占比"], citation_rows)

    doc.add_paragraph()
    doc.add_paragraph(
        "分析结果显示，[数字]格式是主流引用方式（79.6%），这与Water Research、"
        "Science of the Total Environment等主流期刊的格式要求一致。"
        "Author (Year)格式占20.4%，常见于Environmental Science & Technology等期刊。"
    )

    doc.add_page_break()

    # ============================================================
    # 第三章：数据分析方法汇总
    # ============================================================
    doc.add_heading("三、数据分析方法汇总", level=1)

    methods_frequency = analysis_methods.get("methods_frequency", {})

    doc.add_heading("3.1 统计检验方法", level=2)
    doc.add_paragraph(
        "统计检验是环境科学研究中验证假设和判断显著性的核心工具。"
        "以下统计了117篇文献中各类统计检验方法的使用情况。"
    )

    stat_tests = methods_frequency.get("statistical_tests", {})
    st_sorted = sorted(stat_tests.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-1 统计检验方法使用统计")
    add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                     [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in st_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "ANOVA（方差分析）是最常用的统计检验方法（16.2%），用于比较多组数据的差异。"
        "Mann-Whitney U和Kruskal-Wallis等非参数检验方法也有一定使用，"
        "适用于数据不满足正态分布假设的情况。"
    )

    doc.add_heading("3.2 回归分析方法", level=2)
    doc.add_paragraph(
        "回归分析用于探索变量之间的关系和建立预测模型。"
    )

    regression = methods_frequency.get("regression_methods", {})
    reg_sorted = sorted(regression.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-2 回归分析方法使用统计")
    add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                     [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in reg_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "线性回归是最基础且使用最广泛的回归方法（64.1%），用于探索变量间的线性关系。"
        "值得注意的是，正则化方法（Ridge/Lasso）的使用率达到46.2%，"
        "这反映了处理多重共线性和高维数据的需求。"
    )

    doc.add_heading("3.3 机器学习方法", level=2)
    doc.add_paragraph(
        "近年来，机器学习方法在环境科学领域的应用日益广泛，"
        "用于处理复杂的非线性关系和大规模数据。"
    )

    ml_methods = methods_frequency.get("machine_learning", {})
    ml_sorted = sorted(ml_methods.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-3 机器学习方法使用统计")
    add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                     [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in ml_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "随机森林（99.1%）和神经网络（95.7%）是最常用的机器学习方法，"
        "这可能与关键词匹配有关，实际使用率可能较低。"
        "PCA（主成分分析）使用率为23.1%，常用于数据降维和特征提取。"
    )

    doc.add_heading("3.4 不确定性分析方法", level=2)
    doc.add_paragraph(
        "不确定性分析是温室气体排放核算中的重要环节，"
        "用于评估估算结果的可靠性和置信区间。"
    )

    uncertainty = methods_frequency.get("uncertainty_methods", {})
    unc_sorted = sorted(uncertainty.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-4 不确定性分析方法使用统计")
    add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                     [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in unc_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "置信区间（100%）是最基础的不确定性表达方式。"
        "敏感性分析（38.5%）和蒙特卡洛模拟（18.8%）是常用的不确定性量化方法，"
        "用于评估输入参数对输出结果的影响。"
    )

    doc.add_heading("3.5 排放核算方法", level=2)
    doc.add_paragraph(
        "温室气体排放核算方法是本领域的核心内容，"
        "不同方法的适用场景和精度存在差异。"
    )

    emission = methods_frequency.get("emission_accounting", {})
    em_sorted = sorted(emission.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-5 排放核算方法使用统计")
    add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                     [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in em_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "IPCC排放因子法（97.4%）是最广泛使用的核算方法，这与其国际通用性和简便性有关。"
        "生命周期评价（LCA，70.1%）和碳足迹分析（53.8%）也得到广泛应用，"
        "用于评估污水处理全过程的环境影响。"
    )

    doc.add_heading("3.6 数据处理方法", level=2)

    data_proc = methods_frequency.get("data_processing", {})
    dp_sorted = sorted(data_proc.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-6 数据处理方法使用统计")
    add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                     [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in dp_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "数据标准化（24.8%）和异常值检测（13.7%）是常用的数据预处理方法。"
        "对数转换（5.1%）常用于处理右偏分布的数据。"
    )

    doc.add_page_break()

    # ============================================================
    # 第四章：图表使用分析
    # ============================================================
    doc.add_heading("四、图表使用分析", level=1)

    doc.add_heading("4.1 常用图表类型", level=2)
    doc.add_paragraph(
        "图表是数据可视化的重要工具，选择合适的图表类型能够更有效地传达研究结果。"
    )

    figure_types = figure_design.get("figure_types", {})
    ft_sorted = sorted(figure_types.items(), key=lambda x: -x[1])

    doc.add_paragraph("表4-1 常用图表类型统计")
    add_styled_table(doc, ["图表类型", "使用论文数", "使用率"],
                     [[ft, c, f"{c/len(all_results)*100:.1f}%"] for ft, c in ft_sorted])

    doc.add_paragraph()
    doc.add_paragraph(
        "箱线图（box plot）是最常用的图表类型（9.4%），用于展示数据分布和异常值。"
        "热力图（heatmap）和饼图（pie chart）也有一定使用，"
        "分别用于展示相关性矩阵和比例分布。"
    )

    doc.add_heading("4.2 图表数量统计", level=2)

    fig_counts = [r.get("figure_info", {}).get("figure_count", 0) for r in all_results]
    tab_counts = [r.get("figure_info", {}).get("table_count", 0) for r in all_results]

    avg_fig = sum(fig_counts) / len(fig_counts)
    avg_tab = sum(tab_counts) / len(tab_counts)
    max_fig = max(fig_counts)
    max_tab = max(tab_counts)

    doc.add_paragraph("表4-2 图表数量统计")
    add_styled_table(doc, ["统计指标", "图片", "表格"],
                     [["平均数量", f"{avg_fig:.1f} 张/篇", f"{avg_tab:.1f} 个/篇"],
                      ["最大数量", f"{max_fig} 张", f"{max_tab} 个"],
                      ["总计", f"{sum(fig_counts)} 张", f"{sum(tab_counts)} 个"]])

    doc.add_paragraph()
    doc.add_paragraph(
        f"平均每篇论文包含 {avg_fig:.1f} 张图片和 {avg_tab:.1f} 个表格。"
        "这表明图表在数据呈现中占有重要地位。"
    )

    doc.add_heading("4.3 表格复杂度分析", level=2)

    complex_tables = 0
    total_tables = 0
    for r in all_results:
        for tc in r.get("figure_info", {}).get("table_complexity", []):
            total_tables += 1
            if tc.get("has_statistics"):
                complex_tables += 1

    doc.add_paragraph(
        f"在提取的 {total_tables} 个表格中，有 {complex_tables} 个（{complex_tables/total_tables*100:.1f}%）"
        "包含统计信息（如均值、标准差、p值等）。这表明统计结果的表格呈现是常见的做法。"
    )

    doc.add_paragraph()
    doc.add_paragraph("图表使用建议：")
    doc.add_paragraph("使用箱线图展示数据分布和组间比较；", style='List Bullet')
    doc.add_paragraph("使用热力图展示变量间的相关性；", style='List Bullet')
    doc.add_paragraph("使用柱状图比较不同组别的均值；", style='List Bullet')
    doc.add_paragraph("使用散点图展示两个变量之间的关系；", style='List Bullet')
    doc.add_paragraph("表格中应包含均值±标准差、样本量和显著性水平。", style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # 第五章：写作经验总结
    # ============================================================
    doc.add_heading("五、写作经验总结", level=1)

    doc.add_heading("5.1 摘要写作要点", level=2)
    doc.add_paragraph(
        '摘要是论文的"门面"，直接影响读者是否继续阅读全文。'
        '基于文献分析，我们总结以下摘要写作要点：'
    )
    doc.add_paragraph("长度控制在150-300词之间；", style='List Bullet')
    doc.add_paragraph('采用"背景-目的-方法-结果-结论"的结构；', style='List Bullet')
    doc.add_paragraph("结果部分要包含具体数据（数值、p值、相关系数）；", style='List Bullet')
    doc.add_paragraph("避免引用文献和非通用缩写；", style='List Bullet')
    doc.add_paragraph("最后一句应总结研究的主要贡献或意义。", style='List Bullet')

    doc.add_heading("5.2 引言写作模式", level=2)
    doc.add_paragraph(
        '引言部分采用"漏斗结构"，从大背景逐步聚焦到具体研究问题：'
    )
    doc.add_paragraph("第一段：大背景介绍（全球气候变化、污水处理厂排放问题）；", style='List Number')
    doc.add_paragraph("第二段：已有研究综述（排放核算方法、影响因素）；", style='List Number')
    doc.add_paragraph("第三段：研究空白和不足（方法比较、区域差异等）；", style='List Number')
    doc.add_paragraph("第四段：本研究目的和创新点。", style='List Number')

    doc.add_paragraph()
    doc.add_paragraph("引言写作技巧：")
    doc.add_paragraph("每段引用3-5篇文献，支撑论点；", style='List Bullet')
    doc.add_paragraph("使用\"however\"、\"but\"等转折词指出已有研究的不足；", style='List Bullet')
    doc.add_paragraph("使用\"therefore\"、\"thus\"等因果词引出本研究；", style='List Bullet')
    doc.add_paragraph("最后一句明确说明本研究的目的。", style='List Bullet')

    doc.add_heading("5.3 方法描述规范", level=2)
    doc.add_paragraph(
        "方法部分应详细到其他研究者能够重复实验。常见结构包括："
    )
    doc.add_paragraph("采样/实验设计：采样点位、采样频率、样本量；", style='List Bullet')
    doc.add_paragraph("分析方法：仪器型号、分析标准、检测限；", style='List Bullet')
    doc.add_paragraph("统计方法：使用的检验方法、显著性水平、软件版本；", style='List Bullet')
    doc.add_paragraph("核算方法：排放因子来源、计算公式、参数取值。", style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("方法描述示例句式：")
    doc.add_paragraph("\"Greenhouse gas emissions were calculated using the IPCC methodology (IPCC, 2019).\"")
    doc.add_paragraph("\"Statistical analyses were performed using SPSS 26.0 (IBM, USA).\"")
    doc.add_paragraph("\"The significance level was set at p < 0.05.\"")

    doc.add_heading("5.4 结果呈现技巧", level=2)
    doc.add_paragraph(
        "结果部分应客观呈现数据，避免过度解读。常见技巧包括："
    )
    doc.add_paragraph("先文字描述，再引用图表；", style='List Bullet')
    doc.add_paragraph("报告格式：均值 ± 标准差 (n=X)；", style='List Bullet')
    doc.add_paragraph("统计结果格式：F(df1, df2) = X.XX, p = 0.XX；", style='List Bullet')
    doc.add_paragraph("效应量报告：Cohen's d, η², R²；", style='List Bullet')
    doc.add_paragraph("使用\"significantly\"、\"indicated\"等词引导结果。", style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("结果描述示例句式：")
    doc.add_paragraph("\"The results showed that CH4 emissions were significantly higher in summer than in winter (p < 0.01).\"")
    doc.add_paragraph("\"A strong positive correlation was observed between temperature and N2O emissions (r = 0.72, p < 0.001).\"")
    doc.add_paragraph("\"The IPCC method overestimated emissions by a factor of 3.4 compared to direct measurements.\"")

    doc.add_heading("5.5 讨论写作框架", level=2)
    doc.add_paragraph(
        "讨论部分是论文的核心，应将结果与已有研究联系起来。常见框架包括："
    )
    doc.add_paragraph("第一段：总结主要发现；", style='List Number')
    doc.add_paragraph("第二段：与已有研究对比（一致/不一致 + 原因分析）；", style='List Number')
    doc.add_paragraph("第三段：机制解释（为什么会出现这个结果）；", style='List Number')
    doc.add_paragraph("第四段：研究局限性（2-3点）；", style='List Number')
    doc.add_paragraph("第五段：实际意义/政策建议。", style='List Number')

    doc.add_paragraph()
    doc.add_paragraph("讨论写作技巧：")
    doc.add_paragraph("使用\"consistent with\"、\"in agreement with\"表示与已有研究一致；", style='List Bullet')
    doc.add_paragraph("使用\"in contrast\"、\"however\"表示与已有研究不同；", style='List Bullet')
    doc.add_paragraph("使用\"may be attributed to\"、\"could be explained by\"进行机制解释；", style='List Bullet')
    doc.add_paragraph("使用\"limitation\"、\"constraint\"客观指出研究不足。", style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # 第六章：文献列表
    # ============================================================
    doc.add_heading("六、分析文献列表", level=1)
    doc.add_paragraph(
        f"以下为本次分析的 {len(all_results)} 篇文献列表："
    )

    # 分批显示文献列表
    doc.add_paragraph("表6-1 文献列表（第1-40篇）")
    lit_rows_1 = []
    for r in all_results[:40]:
        meta = r.get("metadata", {})
        title = meta.get("title", "")[:40]
        year = meta.get("year", "N/A")
        lit_rows_1.append([r["index"], meta.get("filename", "")[:25], title, year])

    add_styled_table(doc, ["序号", "文件名", "标题", "年份"], lit_rows_1)

    doc.add_paragraph()
    doc.add_paragraph("表6-2 文献列表（第41-80篇）")
    lit_rows_2 = []
    for r in all_results[40:80]:
        meta = r.get("metadata", {})
        title = meta.get("title", "")[:40]
        year = meta.get("year", "N/A")
        lit_rows_2.append([r["index"], meta.get("filename", "")[:25], title, year])

    add_styled_table(doc, ["序号", "文件名", "标题", "年份"], lit_rows_2)

    doc.add_paragraph()
    doc.add_paragraph("表6-3 文献列表（第81-117篇）")
    lit_rows_3 = []
    for r in all_results[80:]:
        meta = r.get("metadata", {})
        title = meta.get("title", "")[:40]
        year = meta.get("year", "N/A")
        lit_rows_3.append([r["index"], meta.get("filename", "")[:25], title, year])

    add_styled_table(doc, ["序号", "文件名", "标题", "年份"], lit_rows_3)

    doc.add_page_break()

    # ============================================================
    # 附录
    # ============================================================
    doc.add_heading("附录", level=1)

    doc.add_heading("附录A：数据来源说明", level=2)
    doc.add_paragraph(
        "本报告数据来源于Web of Science等学术数据库，检索关键词包括："
        "\"greenhouse gases\"、\"wastewater treatment plant\"、\"emission factor\"、"
        "\"carbon footprint\"、\"LCA\"等。文献时间范围为2020-2026年。"
    )

    doc.add_heading("附录B：分析工具说明", level=2)
    doc.add_paragraph(
        "本报告使用Python进行数据分析，主要工具包包括："
    )
    doc.add_paragraph("pdfplumber：PDF文本提取；", style='List Bullet')
    doc.add_paragraph("python-docx：Word文档生成；", style='List Bullet')
    doc.add_paragraph("json：数据存储和读取；", style='List Bullet')
    doc.add_paragraph("re：正则表达式文本处理。", style='List Bullet')

    doc.add_heading("附录C：术语表", level=2)
    terms = [
        ("WWTP", "Wastewater Treatment Plant，污水处理厂"),
        ("GHG", "Greenhouse Gas，温室气体"),
        ("IPCC", "Intergovernmental Panel on Climate Change，政府间气候变化专门委员会"),
        ("LCA", "Life Cycle Assessment，生命周期评价"),
        ("ANOVA", "Analysis of Variance，方差分析"),
        ("PCA", "Principal Component Analysis，主成分分析"),
        ("Hedging", "学术模糊语，用于表达谨慎态度"),
        ("Effect Size", "效应量，用于衡量实际差异大小"),
    ]

    doc.add_paragraph("表C-1 术语表")
    add_styled_table(doc, ["术语", "全称/解释"], terms)

    # ============================================================
    # 保存文档
    # ============================================================
    output_path = os.path.join(OUTPUT_DIR, "文献学习综合分析报告.docx")
    doc.save(output_path)
    print(f"报告已保存: {output_path}")

    return output_path


if __name__ == "__main__":
    generate_report()
