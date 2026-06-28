"""
HRMS-NTS文献章节智能总结生成器

对每篇文章每个章节进行真正的概括总结
"""

import json
import os
import sys
import re
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_full_text(pdf_path, max_pages=30):
    """提取PDF全文"""
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            text_parts = []
            for page in pdf.pages[:max_pages]:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts)
    except Exception as e:
        return None


def split_into_sections(text):
    """将文本分割成章节"""
    sections = {}

    patterns = {
        'abstract': [
            r'(?i)\b(?:Abstract|ABSTRACT|摘要)\b',
        ],
        'introduction': [
            r'(?i)\b(?:\d+\.?\s*)?Introduction\b',
            r'(?i)\b(?:\d+\.?\s*)?INTRODUCTION\b',
        ],
        'methods': [
            r'(?i)\b(?:\d+\.?\s*)?(?:Methods?|Materials?\s*(?:and|&)\s*Methods?|Methodology|Experimental|Experimental\s+Section)\b',
            r'(?i)\b(?:\d+\.?\s*)?(?:MATERIALS?\s+AND\s+METHODS?)\b',
        ],
        'results': [
            r'(?i)\b(?:\d+\.?\s*)?(?:Results?(?:\s*(?:and|&)\s*Discussion)?)\b',
            r'(?i)\b(?:\d+\.?\s*)?RESULTS?\b',
        ],
        'discussion': [
            r'(?i)\b(?:\d+\.?\s*)?Discussion\b',
            r'(?i)\b(?:\d+\.?\s*)?DISCUSSION\b',
        ],
        'conclusion': [
            r'(?i)\b(?:\d+\.?\s*)?(?:Conclusions?|Summary|Concluding\s+Remarks)\b',
            r'(?i)\b(?:\d+\.?\s*)?(?:CONCLUSIONS?|SUMMARY)\b',
        ],
    }

    positions = []
    for section_name, pattern_list in patterns.items():
        for pattern in pattern_list:
            for match in re.finditer(pattern, text):
                start = match.start()
                if start == 0 or text[start-1] in '\n\r \t' or text[start-1].isdigit():
                    end_pos = match.end()
                    if end_pos < len(text) and (text[end_pos] in '\n\r: ' or end_pos == len(text)):
                        positions.append((start, section_name))
                        break
            else:
                continue
            break

    positions.sort(key=lambda x: x[0])

    seen = set()
    unique_positions = []
    for pos, name in positions:
        if name not in seen:
            seen.add(name)
            unique_positions.append((pos, name))

    for i, (start, section_name) in enumerate(unique_positions):
        if i + 1 < len(unique_positions):
            end = unique_positions[i + 1][0]
        else:
            ref_match = re.search(r'(?i)\b(?:References?|Bibliography|REFERENCES?)\b', text[start:])
            if ref_match:
                end = start + ref_match.start()
            else:
                end = len(text)

        section_text = text[start:end].strip()
        lines = section_text.split('\n', 1)
        if len(lines) > 1:
            section_text = lines[1].strip()

        sections[section_name] = section_text[:10000]

    return sections


def extract_key_sentences(text, max_sentences=5):
    """提取关键句子"""
    if not text:
        return []

    # 按句子分割
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    if len(sentences) <= max_sentences:
        return sentences

    # 选择包含关键词的句子
    key_patterns = [
        r'(?i)\b(?:aim|objective|purpose|goal|propose|present|develop|investigate|examine|study|analyze|evaluate|assess|demonstrate|show|reveal|find|conclude|suggest|highlight)\b',
        r'(?i)\b(?:method|approach|technique|procedure|workflow|framework|strategy|protocol)\b',
        r'(?i)\b(?:result|finding|observation|outcome|performance|efficiency|accuracy|significant)\b',
        r'(?i)\b(?:novel|new|innovative|first|unique|important|significant|promising)\b',
        r'(?i)\b(?:challenge|limitation|problem|issue|difficulty|gap)\b',
        r'(?i)\b(?:future|direction|recommendation|perspective|outlook)\b',
    ]

    scored_sentences = []
    for i, sentence in enumerate(sentences):
        score = 0
        # 首尾句子权重更高
        if i < 3:
            score += 2
        if i >= len(sentences) - 3:
            score += 1
        # 包含关键词的句子
        for pattern in key_patterns:
            if re.search(pattern, sentence):
                score += 1
        scored_sentences.append((score, i, sentence))

    # 按分数排序，取前N个
    scored_sentences.sort(key=lambda x: -x[0])
    selected = sorted(scored_sentences[:max_sentences], key=lambda x: x[1])

    return [s[2] for s in selected]


def summarize_abstract(text):
    """总结摘要"""
    if not text or len(text.strip()) < 50:
        return "摘要内容未能提取到。"

    # 摘要通常包含：背景、目的、方法、结果、结论
    sentences = extract_key_sentences(text, 8)
    return ' '.join(sentences)


def summarize_introduction(text):
    """总结引言"""
    if not text or len(text.strip()) < 100:
        return "引言内容未能提取到。"

    # 引言通常包含：研究背景、问题陈述、研究目的、文章结构
    key_info = []

    # 提取研究背景
    bg_patterns = [
        r'(?i)(?:background|context|overview|recent|currently|growing|increasing).*?\.',
        r'(?i)(?:pollution|contamination|environment|health|risk|concern).*?\.',
    ]

    # 提取研究目的
    purpose_patterns = [
        r'(?i)(?:aim|objective|purpose|goal|this (?:study|paper|work|research|review)).*?\.',
        r'(?i)(?:propose|present|develop|investigate|examine|focus).*?\.',
    ]

    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    # 找背景句
    background = []
    for s in sentences[:10]:
        for pattern in bg_patterns:
            if re.search(pattern, s):
                background.append(s)
                break

    # 找目的句
    purposes = []
    for s in sentences:
        for pattern in purpose_patterns:
            if re.search(pattern, s):
                purposes.append(s)
                break

    if background:
        key_info.append("【研究背景】" + ' '.join(background[:2]))
    if purposes:
        key_info.append("【研究目的】" + ' '.join(purposes[:2]))

    return ' '.join(key_info) if key_info else ' '.join(sentences[:5])


def summarize_methods(text):
    """总结方法"""
    if not text or len(text.strip()) < 100:
        return "方法内容未能提取到。"

    key_info = []

    # 提取关键方法信息
    method_keywords = {
        '样品采集': r'(?i)(?:sample|sampling|collect|collection|field)',
        '样品处理': r'(?i)(?:extract|extraction|pretreatment|preparation|clean.?up|SPE|LLE)',
        '分析仪器': r'(?i)(?:LC|GC|MS|HRMS|Orbitrap|QTOF|TOF|FT-ICR|instrument)',
        '数据处理': r'(?i)(?:software|processing|peak|alignment|normalization|XCMS|MZmine|MS-DIAL)',
        '筛查策略': r'(?i)(?:target|suspect|non.?target|screening|identification)',
        '质量控制': r'(?i)(?:quality|control|QA|QC|blank|standard|spike)',
    }

    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    for category, pattern in method_keywords.items():
        for s in sentences[:50]:
            if re.search(pattern, s):
                key_info.append(f"【{category}】{s[:200]}")
                break

    return '\n'.join(key_info) if key_info else ' '.join(sentences[:5])


def summarize_results(text):
    """总结结果"""
    if not text or len(text.strip()) < 100:
        return "结果内容未能提取到。"

    key_info = []

    # 提取关键结果
    result_patterns = [
        r'(?i)(?:identified|detected|found|discovered|observed|revealed|showed|demonstrated).*?\.',
        r'(?i)(?:\d+ (?:compounds|chemicals|features|peaks|analytes|targets|suspects)).*?\.',
        r'(?i)(?:concentration|level|range|mean|average|median|±).*?\.',
        r'(?i)(?:significant|correlation|relationship|association|trend).*?\.',
    ]

    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    # 提取数字相关的发现
    number_findings = []
    for s in sentences:
        if re.search(r'\d+', s) and len(s) < 300:
            for pattern in result_patterns:
                if re.search(pattern, s):
                    number_findings.append(s[:200])
                    break

    if number_findings:
        key_info.append("【主要发现】" + ' '.join(number_findings[:3]))

    # 提取其他重要发现
    other_findings = []
    for s in sentences:
        if re.search(r'(?i)(?:important|significant|notable|remarkable|novel|key)', s):
            other_findings.append(s[:200])
            if len(other_findings) >= 3:
                break

    if other_findings:
        key_info.append("【重要发现】" + ' '.join(other_findings))

    return '\n'.join(key_info) if key_info else ' '.join(sentences[:5])


def summarize_discussion(text):
    """总结讨论"""
    if not text or len(text.strip()) < 100:
        return "讨论内容未能提取到。"

    key_info = []

    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    # 提取比较讨论
    comparison_patterns = [
        r'(?i)(?:compar|higher|lower|greater|less|better|worse|similar|different|consistent|agreement).*?\.',
        r'(?i)(?:literature|previous|reported|studies|other).*?\.',
    ]

    comparisons = []
    for s in sentences:
        for pattern in comparison_patterns:
            if re.search(pattern, s):
                comparisons.append(s[:200])
                break
        if len(comparisons) >= 3:
            break

    if comparisons:
        key_info.append("【与已有研究比较】" + ' '.join(comparisons))

    # 提取机制解释
    mechanism_patterns = [
        r'(?i)(?:mechanism|explain|reason|cause|factor|influence|affect|impact|due to|because).*?\.',
    ]

    mechanisms = []
    for s in sentences:
        for pattern in mechanism_patterns:
            if re.search(pattern, s):
                mechanisms.append(s[:200])
                break
        if len(mechanisms) >= 2:
            break

    if mechanisms:
        key_info.append("【机制解释】" + ' '.join(mechanisms))

    # 提取局限性
    limitation_patterns = [
        r'(?i)(?:limitation|limit|challenge|problem|issue|drawback|weakness|gap|difficulty|uncertain).*?\.',
    ]

    limitations = []
    for s in sentences:
        for pattern in limitation_patterns:
            if re.search(pattern, s):
                limitations.append(s[:200])
                break
        if len(limitations) >= 2:
            break

    if limitations:
        key_info.append("【研究局限】" + ' '.join(limitations))

    return '\n'.join(key_info) if key_info else ' '.join(sentences[:5])


def summarize_conclusion(text):
    """总结结论"""
    if not text or len(text.strip()) < 50:
        return "结论内容未能提取到。"

    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 30]

    key_info = []

    # 提取主要结论
    conclusion_patterns = [
        r'(?i)(?:conclude|conclusion|summary|overall|in summary|in conclusion|this study|our (?:study|results|findings)).*?\.',
        r'(?i)(?:demonstrate|show|reveal|confirm|suggest|indicate|highlight|emphasize).*?\.',
    ]

    conclusions = []
    for s in sentences:
        for pattern in conclusion_patterns:
            if re.search(pattern, s):
                conclusions.append(s[:250])
                break
        if len(conclusions) >= 3:
            break

    if conclusions:
        key_info.append("【主要结论】" + ' '.join(conclusions))

    # 提取未来展望
    future_patterns = [
        r'(?i)(?:future|further|next|prospect|perspective|recommendation|need|should|require).*?\.',
    ]

    futures = []
    for s in sentences:
        for pattern in future_patterns:
            if re.search(pattern, s):
                futures.append(s[:200])
                break
        if len(futures) >= 2:
            break

    if futures:
        key_info.append("【未来展望】" + ' '.join(futures))

    return '\n'.join(key_info) if key_info else ' '.join(sentences[:5])


def main():
    """主函数"""
    json_path = r'd:\VScode\firstcc\GHGs-WWTPs\output\literature_learning\all_papers_analysis.json'
    literature_dir = r'D:\下载\现代环境分析技术'
    output_path = r'd:\VScode\firstcc\GHGs-WWTPs\output\literature_learning\HRMS-NTS文献各章节内容总结.docx'

    print("加载文献列表...")
    with open(json_path, 'r', encoding='utf-8') as f:
        papers_data = json.load(f)
    print(f"共 {len(papers_data)} 篇文献")

    if not HAS_DOCX:
        print("错误：未安装python-docx")
        return

    doc = Document()

    # 标题
    title = doc.add_heading('高分辨率质谱非靶向筛查（HRMS-NTS）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('文献各章节内容总结', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'文献总数：{len(papers_data)} 篇')
    doc.add_paragraph('说明：本报告对每篇文献的各个章节进行了概括总结，提取关键信息。')
    doc.add_paragraph('')

    # 逐篇处理
    for idx, paper_meta in enumerate(papers_data):
        filename = paper_meta.get('metadata', {}).get('filename', '')
        if not filename:
            continue

        pdf_path = os.path.join(literature_dir, filename)
        if not os.path.exists(pdf_path):
            continue

        print(f"[{idx+1}/{len(papers_data)}] 处理: {filename[:40]}...")

        full_text = extract_full_text(pdf_path)
        if not full_text:
            continue

        sections = split_into_sections(full_text)
        meta = paper_meta.get('metadata', {})
        title_text = meta.get('title', filename)[:80]

        # 论文标题
        doc.add_heading(f'文献 {idx+1}：{title_text}', level=1)

        # 基本信息（简化）
        p = doc.add_paragraph()
        p.add_run('文件名：').bold = True
        p.add_run(filename)

        if meta.get('year'):
            p = doc.add_paragraph()
            p.add_run('年份：').bold = True
            p.add_run(str(meta['year']))

        doc.add_paragraph('')

        # 各章节总结
        summaries = {
            '摘要': summarize_abstract(sections.get('abstract', '')),
            '引言': summarize_introduction(sections.get('introduction', '')),
            '方法': summarize_methods(sections.get('methods', '')),
            '结果': summarize_results(sections.get('results', '')),
            '讨论': summarize_discussion(sections.get('discussion', '')),
            '结论': summarize_conclusion(sections.get('conclusion', '')),
        }

        for section_name, summary in summaries.items():
            doc.add_heading(section_name, level=2)
            if summary and len(summary) > 20:
                doc.add_paragraph(summary)
            else:
                doc.add_paragraph('【本章节内容未能有效提取】', style='Intense Quote')

        # 分页
        if idx < len(papers_data) - 1:
            doc.add_page_break()

    print(f"\n保存报告...")
    doc.save(output_path)
    print(f"报告已保存: {output_path}")
    print(f"文件大小: {os.path.getsize(output_path) / 1024:.0f} KB")


if __name__ == '__main__':
    main()
