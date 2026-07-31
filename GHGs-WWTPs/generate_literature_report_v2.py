"""
生成文献学习综合分析报告 v2（Word格式）— 修正版
"""

import os
import sys
import json
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    sys.exit(1)

# ============================================================
# 配置
# ============================================================
DATA_DIR = r"D:\VScode\firstcc\GHGs-WWTPs\output\literature_learning"
KNOWLEDGE_DIR = r"D:\VScode\firstcc\GHGs-WWTPs\knowledge_store"
OUTPUT_DIR = r"D:\VScode\firstcc\GHGs-WWTPs\output\literature_learning"

# ============================================================
# 辅助函数
# ============================================================

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E75B6')

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            # 交替行颜色
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'D6E4F0')

    return table


# ============================================================
# 主报告生成
# ============================================================

def generate_report():
    """生成完整的Word报告"""

    # 加载数据
    with open(os.path.join(DATA_DIR, "all_papers_analysis_v2.json"), "r", encoding="utf-8") as f:
        all_results = json.load(f)

    with open(os.path.join(KNOWLEDGE_DIR, "learned_writing_patterns_v2.json"), "r", encoding="utf-8") as f:
        writing_patterns = json.load(f)

    with open(os.path.join(KNOWLEDGE_DIR, "learned_analysis_methods_v2.json"), "r", encoding="utf-8") as f:
        analysis_methods = json.load(f)

    with open(os.path.join(KNOWLEDGE_DIR, "learned_figure_design_v2.json"), "r", encoding="utf-8") as f:
        figure_design = json.load(f)

    # 创建文档
    doc = Document()

    # ============================================================
    # 封面
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("污水处理厂温室气体排放文献\n综合分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("——写作模式、数据分析方法与图表设计学习总结")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    version_note = doc.add_paragraph()
    version_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_note.add_run("【修正版 v2 — 严格方法识别】")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph()

    # 信息框
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"分析文献数量：{len(all_results)} 篇\n")
    run.font.size = Pt(12)
    run = info_para.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}\n")
    run.font.size = Pt(12)
    run = info_para.add_run("分析工具：GHGs-WWTPs 文献学习系统 v2")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ============================================================
    # 重要说明
    # ============================================================
    doc.add_heading("重要说明", level=1)
    doc.add_paragraph(
        "本报告采用严格的方法识别逻辑，与初版报告有显著差异："
    )
    doc.add_paragraph("")
    doc.add_paragraph("1. 只统计论文中实际使用的方法，排除仅在文献综述或讨论中提到的方法；", style='List Number')
    doc.add_paragraph("2. 方法识别基于上下文语境，需要出现'was used'、'was performed'等使用动词；", style='List Number')
    doc.add_paragraph("3. 优先在Methods部分搜索，确保方法是本文实际采用的；", style='List Number')
    doc.add_paragraph("4. 排除引用他人研究时提到的方法。", style='List Number')

    doc.add_paragraph()
    doc.add_paragraph(
        "因此，本报告中的方法使用率普遍低于初版报告，但数据更加真实可靠。"
    )

    doc.add_page_break()

    # ============================================================
    # 目录页
    # ============================================================
    doc.add_heading("目  录", level=1)
    doc.add_paragraph()

    toc_items = [
        ("重要说明", 1),
        ("一、研究概述", 3),
        ("    1.1 研究背景与目的", 3),
        ("    1.2 文献来源与范围", 3),
        ("二、写作模式分析", 4),
        ("    2.1 高频过渡词分析", 4),
        ("    2.2 学术模糊语使用", 5),
        ("    2.3 强调语使用分析", 5),
        ("    2.4 引用格式统计", 6),
        ("三、数据分析方法汇总（修正版）", 7),
        ("    3.1 统计检验方法", 7),
        ("    3.2 回归分析方法", 7),
        ("    3.3 机器学习方法", 8),
        ("    3.4 不确定性分析方法", 8),
        ("    3.5 排放核算方法", 9),
        ("    3.6 数据处理方法", 9),
        ("四、图表使用分析", 10),
        ("五、写作经验总结", 11),
        ("六、文献列表", 12),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item} {'.' * (50 - len(item))} {page}")
        run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # 第一章：研究概述
    # ============================================================
    doc.add_heading("一、研究概述", level=1)

    doc.add_heading("1.1 研究背景与目的", level=2)
    doc.add_paragraph(
        "污水处理厂（WWTPs）是温室气体（GHG）排放的重要来源之一，主要包括二氧化碳（CO2）、"
        "甲烷（CH4）和氧化亚氮（N2O）。准确核算和有效控制污水处理过程中的温室气体排放，"
        "对于实现碳中和目标具有重要意义。"
    )
    doc.add_paragraph(
        "本报告通过对117篇相关文献进行系统分析，旨在："
    )
    doc.add_paragraph("（1）总结文献中常用的写作模式和学术表达方式；", style='List Bullet')
    doc.add_paragraph("（2）梳理文献中实际使用的数据分析方法及其应用频率；", style='List Bullet')
    doc.add_paragraph("（3）归纳图表设计和数据可视化的方法与技巧；", style='List Bullet')
    doc.add_paragraph("（4）为后续论文写作提供经验参考和方法指导。", style='List Bullet')

    doc.add_heading("1.2 文献来源与范围", level=2)
    doc.add_paragraph(
        f"本次分析共收录 {len(all_results)} 篇文献，主要来源于Web of Science等学术数据库，"
        "涵盖2020-2026年间发表的高水平期刊论文。文献主题涵盖污水处理厂温室气体排放的核算方法、"
        "影响因素、减排技术等多个方面。"
    )

    # 文献年份分布
    year_counts = {}
    for r in all_results:
        year = r.get("metadata", {}).get("year", "未知")
        if year and year != "未知":
            try:
                year = int(year)
                if 2000 <= year <= 2030:
                    year_counts[year] = year_counts.get(year, 0) + 1
            except:
                pass

    doc.add_paragraph("文献年份分布如下表所示：")
    year_rows = sorted(year_counts.items(), key=lambda x: x[0])
    add_styled_table(doc, ["年份", "文献数量", "占比"],
                     [[y, c, f"{c/len(all_results)*100:.1f}%"] for y, c in year_rows])

    doc.add_page_break()

    # ============================================================
    # 第二章：写作模式分析
    # ============================================================
    doc.add_heading("二、写作模式分析", level=1)

    doc.add_heading("2.1 高频过渡词分析", level=2)
    doc.add_paragraph(
        "过渡词是学术写作中连接句子和段落的重要工具。通过对117篇文献的分析，"
        "我们统计了各类过渡词的使用频率，为后续写作提供参考。"
    )

    # 过渡词表格
    transition_words = writing_patterns.get("transition_words", {})
    tw_sorted = sorted(transition_words.items(), key=lambda x: -x[1])[:20]

    doc.add_paragraph("表2-1 高频过渡词统计（Top 20）")
    add_styled_table(doc, ["排名", "过渡词", "出现总次数", "平均频率（次/篇）"],
                     [[i+1, w, c, f"{c/len(all_results):.1f}"] for i, (w, c) in enumerate(tw_sorted)])

    doc.add_paragraph()
    doc.add_paragraph(
        '分析结果显示，"however"是使用频率最高的过渡词（平均5.7次/篇），'
        '表明学术写作中对比和转折是常见的逻辑关系。"therefore"和"thus"的高频率'
        '说明因果关系的表达在论文中同样重要。'
    )

    doc.add_heading("2.2 学术模糊语（Hedging）使用", level=2)
    doc.add_paragraph(
        "学术模糊语是学术写作中的重要特征，用于表达谨慎的态度和避免过于绝对的结论。"
    )

    hedging_phrases = writing_patterns.get("hedging_phrases", {})
    hp_sorted = sorted(hedging_phrases.items(), key=lambda x: -x[1])[:15]

    doc.add_paragraph("表2-2 学术模糊语使用频率")
    add_styled_table(doc, ["模糊语", "出现次数", "平均频率（次/篇）"],
                     [[w, c, f"{c/len(all_results):.1f}"] for w, c in hp_sorted])

    doc.add_heading("2.3 强调语使用分析", level=2)

    emphasis_phrases = writing_patterns.get("emphasis_phrases", {})
    ep_sorted = sorted(emphasis_phrases.items(), key=lambda x: -x[1])[:10]

    doc.add_paragraph("表2-3 强调语使用频率")
    add_styled_table(doc, ["强调语", "出现次数", "平均频率（次/篇）"],
                     [[w, c, f"{c/len(all_results):.1f}"] for w, c in ep_sorted])

    doc.add_heading("2.4 引用格式统计", level=2)

    citation_patterns = writing_patterns.get("citation_patterns", [])
    citation_summary = {}
    for cp in citation_patterns:
        pattern = cp.get("pattern", "")
        count = cp.get("count", 0)
        if "A-Z" in pattern or "Author" in pattern:
            citation_summary["Author (Year)"] = citation_summary.get("Author (Year)", 0) + count
        elif "\\d" in pattern:
            citation_summary["[Number]"] = citation_summary.get("[Number]", 0) + count

    doc.add_paragraph("表2-4 引用格式统计")
    citation_rows = [[fmt, count, f"{count/sum(citation_summary.values())*100:.1f}%"]
                     for fmt, count in sorted(citation_summary.items(), key=lambda x: -x[1])]
    add_styled_table(doc, ["引用格式", "使用次数", "占比"], citation_rows)

    doc.add_page_break()

    # ============================================================
    # 第三章：数据分析方法汇总（修正版）
    # ============================================================
    doc.add_heading("三、数据分析方法汇总（修正版）", level=1)

    doc.add_paragraph(
        "【重要】本章数据采用严格的方法识别逻辑，只统计论文中实际使用的方法，"
        "排除了仅在文献综述或讨论中提到的方法。因此使用率普遍低于初版报告，但更加真实可靠。"
    )

    methods_frequency = analysis_methods.get("methods_frequency", {})

    # 3.1 统计检验
    doc.add_heading("3.1 统计检验方法", level=2)
    doc.add_paragraph(
        "统计检验是环境科学研究中验证假设和判断显著性的核心工具。"
    )

    stat_tests = methods_frequency.get("statistical_tests", {})
    st_sorted = sorted(stat_tests.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-1 统计检验方法使用统计")
    if st_sorted:
        add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                         [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in st_sorted])
    else:
        doc.add_paragraph("未检测到明确使用统计检验方法的论文。")

    doc.add_paragraph()
    doc.add_paragraph(
        "ANOVA（方差分析）是使用最多的统计检验方法（2.6%），用于比较多组数据的差异。"
        "Shapiro-Wilk和Levene检验用于数据正态性和方差齐性检验。"
    )

    # 3.2 回归分析
    doc.add_heading("3.2 回归分析方法", level=2)

    regression = methods_frequency.get("regression_methods", {})
    reg_sorted = sorted(regression.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-2 回归分析方法使用统计")
    if reg_sorted:
        add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                         [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in reg_sorted])
    else:
        doc.add_paragraph("未检测到明确使用回归分析方法的论文。")

    # 3.3 机器学习
    doc.add_heading("3.3 机器学习方法", level=2)
    doc.add_paragraph(
        "近年来，机器学习方法在环境科学领域的应用日益广泛。"
        "以下统计仅包含论文中实际构建和训练的模型，排除了仅在讨论中提到的方法。"
    )

    ml_methods = methods_frequency.get("machine_learning", {})
    ml_sorted = sorted(ml_methods.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-3 机器学习方法使用统计")
    if ml_sorted:
        add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                         [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in ml_sorted])
    else:
        doc.add_paragraph("未检测到明确使用机器学习方法的论文。")

    doc.add_paragraph()
    doc.add_paragraph(
        "神经网络（23.9%）是最常用的机器学习方法，主要用于排放预测和建模。"
        "随机森林（6.0%）和PCA（3.4%）也有一定应用。"
        "与初版报告相比，使用率大幅下降，但数据更加真实。"
    )

    # 3.4 不确定性分析
    doc.add_heading("3.4 不确定性分析方法", level=2)

    uncertainty = methods_frequency.get("uncertainty_methods", {})
    unc_sorted = sorted(uncertainty.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-4 不确定性分析方法使用统计")
    if unc_sorted:
        add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                         [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in unc_sorted])
    else:
        doc.add_paragraph("未检测到明确使用不确定性分析方法的论文。")

    # 3.5 排放核算
    doc.add_heading("3.5 排放核算方法", level=2)
    doc.add_paragraph(
        "温室气体排放核算方法是本领域的核心内容。"
    )

    emission = methods_frequency.get("emission_accounting", {})
    em_sorted = sorted(emission.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-5 排放核算方法使用统计")
    if em_sorted:
        add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                         [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in em_sorted])
    else:
        doc.add_paragraph("未检测到明确使用排放核算方法的论文。")

    doc.add_paragraph()
    doc.add_paragraph(
        "IPCC排放因子法（43.6%）是最广泛使用的核算方法，这与其国际通用性和简便性有关。"
        "生命周期评价（LCA，11.1%）和碳足迹分析（7.7%）也得到一定应用。"
    )

    # 3.6 数据处理
    doc.add_heading("3.6 数据处理方法", level=2)

    data_proc = methods_frequency.get("data_processing", {})
    dp_sorted = sorted(data_proc.items(), key=lambda x: -x[1])

    doc.add_paragraph("表3-6 数据处理方法使用统计")
    if dp_sorted:
        add_styled_table(doc, ["方法", "使用论文数", "使用率"],
                         [[m, c, f"{c/len(all_results)*100:.1f}%"] for m, c in dp_sorted])
    else:
        doc.add_paragraph("未检测到明确使用数据处理方法的论文。")

    doc.add_page_break()

    # ============================================================
    # 第四章：图表使用分析
    # ============================================================
    doc.add_heading("四、图表使用分析", level=1)

    doc.add_heading("4.1 常用图表类型", level=2)

    figure_types = figure_design.get("figure_types", {})
    ft_sorted = sorted(figure_types.items(), key=lambda x: -x[1])

    doc.add_paragraph("表4-1 常用图表类型统计")
    if ft_sorted:
        add_styled_table(doc, ["图表类型", "使用论文数", "使用率"],
                         [[ft, c, f"{c/len(all_results)*100:.1f}%"] for ft, c in ft_sorted])
    else:
        doc.add_paragraph("未检测到明确的图表类型。")

    doc.add_heading("4.2 图表数量统计", level=2)

    fig_counts = [r.get("figure_info", {}).get("figure_count", 0) for r in all_results]
    tab_counts = [r.get("figure_info", {}).get("table_count", 0) for r in all_results]

    avg_fig = sum(fig_counts) / len(fig_counts) if fig_counts else 0
    avg_tab = sum(tab_counts) / len(tab_counts) if tab_counts else 0
    max_fig = max(fig_counts) if fig_counts else 0
    max_tab = max(tab_counts) if tab_counts else 0

    doc.add_paragraph("表4-2 图表数量统计")
    add_styled_table(doc, ["统计指标", "图片", "表格"],
                     [["平均数量", f"{avg_fig:.1f} 张/篇", f"{avg_tab:.1f} 个/篇"],
                      ["最大数量", f"{max_fig} 张", f"{max_tab} 个"],
                      ["总计", f"{sum(fig_counts)} 张", f"{sum(tab_counts)} 个"]])

    doc.add_page_break()

    # ============================================================
    # 第五章：写作经验总结
    # ============================================================
    doc.add_heading("五、写作经验总结", level=1)

    doc.add_heading("5.1 摘要写作要点", level=2)
    doc.add_paragraph("长度控制在150-300词之间；", style='List Bullet')
    doc.add_paragraph('采用"背景-目的-方法-结果-结论"的结构；', style='List Bullet')
    doc.add_paragraph("结果部分要包含具体数据（数值、p值、相关系数）；", style='List Bullet')
    doc.add_paragraph("避免引用文献和非通用缩写；", style='List Bullet')
    doc.add_paragraph("最后一句应总结研究的主要贡献或意义。", style='List Bullet')

    doc.add_heading("5.2 引言写作模式", level=2)
    doc.add_paragraph('引言部分采用"漏斗结构"，从大背景逐步聚焦到具体研究问题：')
    doc.add_paragraph("第一段：大背景介绍（全球气候变化、污水处理厂排放问题）；", style='List Number')
    doc.add_paragraph("第二段：已有研究综述（排放核算方法、影响因素）；", style='List Number')
    doc.add_paragraph("第三段：研究空白和不足（方法比较、区域差异等）；", style='List Number')
    doc.add_paragraph("第四段：本研究目的和创新点。", style='List Number')

    doc.add_heading("5.3 方法描述规范", level=2)
    doc.add_paragraph("采样/实验设计：采样点位、采样频率、样本量；", style='List Bullet')
    doc.add_paragraph("分析方法：仪器型号、分析标准、检测限；", style='List Bullet')
    doc.add_paragraph("统计方法：使用的检验方法、显著性水平、软件版本；", style='List Bullet')
    doc.add_paragraph("核算方法：排放因子来源、计算公式、参数取值。", style='List Bullet')

    doc.add_heading("5.4 结果呈现技巧", level=2)
    doc.add_paragraph("先文字描述，再引用图表；", style='List Bullet')
    doc.add_paragraph("报告格式：均值 ± 标准差 (n=X)；", style='List Bullet')
    doc.add_paragraph("统计结果格式：F(df1, df2) = X.XX, p = 0.XX；", style='List Bullet')
    doc.add_paragraph("效应量报告：Cohen's d, η², R²。", style='List Bullet')

    doc.add_heading("5.5 讨论写作框架", level=2)
    doc.add_paragraph("第一段：总结主要发现；", style='List Number')
    doc.add_paragraph("第二段：与已有研究对比（一致/不一致 + 原因分析）；", style='List Number')
    doc.add_paragraph("第三段：机制解释（为什么会出现这个结果）；", style='List Number')
    doc.add_paragraph("第四段：研究局限性（2-3点）；", style='List Number')
    doc.add_paragraph("第五段：实际意义/政策建议。", style='List Number')

    doc.add_page_break()

    # ============================================================
    # 第六章：文献列表
    # ============================================================
    doc.add_heading("六、分析文献列表", level=1)
    doc.add_paragraph(
        f"以下为本次分析的 {len(all_results)} 篇文献列表："
    )

    # 分批显示文献列表
    doc.add_paragraph("表6-1 文献列表（第1-40篇）")
    lit_rows_1 = []
    for r in all_results[:40]:
        meta = r.get("metadata", {})
        title = meta.get("title", "")[:40]
        year = meta.get("year", "N/A")
        lit_rows_1.append([r["index"], meta.get("filename", "")[:25], title, year])

    add_styled_table(doc, ["序号", "文件名", "标题", "年份"], lit_rows_1)

    doc.add_paragraph()
    doc.add_paragraph("表6-2 文献列表（第41-80篇）")
    lit_rows_2 = []
    for r in all_results[40:80]:
        meta = r.get("metadata", {})
        title = meta.get("title", "")[:40]
        year = meta.get("year", "N/A")
        lit_rows_2.append([r["index"], meta.get("filename", "")[:25], title, year])

    add_styled_table(doc, ["序号", "文件名", "标题", "年份"], lit_rows_2)

    doc.add_paragraph()
    doc.add_paragraph("表6-3 文献列表（第81-117篇）")
    lit_rows_3 = []
    for r in all_results[80:]:
        meta = r.get("metadata", {})
        title = meta.get("title", "")[:40]
        year = meta.get("year", "N/A")
        lit_rows_3.append([r["index"], meta.get("filename", "")[:25], title, year])

    add_styled_table(doc, ["序号", "文件名", "标题", "年份"], lit_rows_3)

    # ============================================================
    # 保存文档
    # ============================================================
    output_path = os.path.join(OUTPUT_DIR, "文献学习综合分析报告_v3.docx")
    doc.save(output_path)
    print(f"报告已保存: {output_path}")

    return output_path


if __name__ == "__main__":
    generate_report()
