"""
HRMS-NTS文献统计报告生成器

基于文献阅读系统的结果，生成详细的Word报告
"""

import json
import os
import sys
from datetime import datetime
from collections import Counter, defaultdict

sys.stdout.reconfigure(encoding='utf-8')

# 检查python-docx是否可用
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: 未安装python-docx，将生成Markdown报告")


def load_analysis_data(json_path):
    """加载分析数据"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def extract_statistics(data):
    """提取统计数据"""
    stats = {
        'total_papers': len(data),
        'year_distribution': Counter(),
        'journal_distribution': Counter(),
        'region_distribution': Counter(),

        # 分析平台
        'platform': {
            'Orbitrap': 0,
            'QTOF': 0,
            'TOF': 0,
            'FT-ICR': 0,
        },
        'chromatography': {
            'LC-HRMS': 0,
            'GC-HRMS': 0,
        },

        # 样品类型
        'sample_type': {
            'surface water': 0,
            'groundwater': 0,
            'drinking water': 0,
            'wastewater': 0,
            'soil': 0,
            'biota': 0,
        },

        # 筛查策略
        'screening_strategy': {
            'target analysis': 0,
            'suspect screening': 0,
            'non-target screening': 0,
            'combined strategy': 0,
        },

        # 数据处理软件
        'data_software': {
            'Compound Discoverer': 0,
            'MS-DIAL': 0,
            'MZmine': 0,
            'XCMS': 0,
        },

        # 数据库
        'database': {
            'MassBank': 0,
            'mzCloud': 0,
            'NORMAN SusDat': 0,
            'EPA CompTox': 0,
            'PubChem': 0,
        },

        # 污染物类型
        'pollutant_type': {
            'PFAS': 0,
            'pharmaceuticals': 0,
            'pesticides': 0,
            'endocrine disruptors': 0,
            'flame retardants': 0,
            'industrial chemicals': 0,
        },

        # 鉴定等级
        'identification_level': {
            'Schymanski Level 1': 0,
            'Schymanski Level 2': 0,
            'Schymanski Level 3': 0,
            'Schymanski Level 4': 0,
            'Schymanski Level 5': 0,
        },

        # 技术优势和问题
        'advantages': [],
        'problems': [],
        'trends': [],
    }

    for paper in data:
        meta = paper.get('metadata', {})
        methods = paper.get('analysis_methods', {})

        # 年份分布
        year = meta.get('year', 'Unknown')
        if year:
            stats['year_distribution'][year] += 1

        # 期刊分布 - 从标题或文件名提取
        title = meta.get('title', '')
        filename = meta.get('filename', '')

        # 从文件名提取期刊信息
        if 'S00219673' in filename:
            stats['journal_distribution']['Journal of Chromatography A'] += 1
        elif 'S00399140' in filename:
            stats['journal_distribution']['Talanta'] += 1
        elif 'S00431354' in filename:
            stats['journal_distribution']['Water Research'] += 1
        elif 'S00489697' in filename:
            stats['journal_distribution']['Science of the Total Environment'] += 1
        elif 'S01476513' in filename:
            stats['journal_distribution']['Ecotoxicology and Environmental Safety'] += 1
        elif 'S01604120' in filename:
            stats['journal_distribution']['Environment International'] += 1
        elif 'S01659936' in filename:
            stats['journal_distribution']['Trends in Environmental Analytical Chemistry'] += 1
        elif 'S22141588' in filename:
            stats['journal_distribution']['Journal of Pharmaceutical and Biomedical Analysis'] += 1
        else:
            stats['journal_distribution']['Other'] += 1

        # 分析平台
        platform_list = methods.get('analysis_platform', [])
        for platform in ['Orbitrap', 'QTOF', 'TOF', 'FT-ICR']:
            if platform in platform_list:
                stats['platform'][platform] += 1

        # 色谱技术
        for chrom in ['LC-HRMS', 'GC-HRMS']:
            if chrom in platform_list:
                stats['chromatography'][chrom] += 1

        # 样品类型
        sample_list = methods.get('sample_type', [])
        for sample in ['surface water', 'groundwater', 'drinking water', 'wastewater', 'soil', 'biota']:
            if sample in sample_list:
                stats['sample_type'][sample] += 1

        # 筛查策略
        strategy_list = methods.get('screening_strategy', [])
        for strategy in ['target analysis', 'suspect screening', 'non-target screening', 'combined strategy']:
            if strategy in strategy_list:
                stats['screening_strategy'][strategy] += 1

        # 数据处理软件
        software_list = methods.get('data_processing', [])
        for software in ['Compound Discoverer', 'MS-DIAL', 'MZmine', 'XCMS']:
            if software in software_list:
                stats['data_software'][software] += 1

        # 数据库
        for db in ['MassBank', 'mzCloud', 'NORMAN SusDat', 'EPA CompTox', 'PubChem']:
            if db in software_list:
                stats['database'][db] += 1

        # 污染物类型
        pollutant_list = methods.get('pollutant_type', [])
        for pollutant in ['PFAS', 'pharmaceuticals', 'pesticides', 'endocrine disruptors', 'flame retardants', 'industrial chemicals']:
            if pollutant in pollutant_list:
                stats['pollutant_type'][pollutant] += 1

        # 鉴定等级
        level_list = methods.get('identification_level', [])
        for level in ['Schymanski Level 1', 'Schymanski Level 2', 'Schymanski Level 3', 'Schymanski Level 4', 'Schymanski Level 5']:
            if level in level_list:
                stats['identification_level'][level] += 1

    return stats


def generate_markdown_report(data, stats, output_path):
    """生成Markdown报告"""
    report = []
    report.append("# 高分辨率质谱非靶向筛查（HRMS-NTS）文献统计报告\n")
    report.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    report.append(f"**文献总数**: {stats['total_papers']} 篇\n")

    # 1. 文献基本信息
    report.append("\n## 1. 文献基本信息\n")

    report.append("### 1.1 发表年份分布\n")
    report.append("| 年份 | 文献数量 | 占比 |")
    report.append("|------|----------|------|")
    for year in sorted(stats['year_distribution'].keys()):
        count = stats['year_distribution'][year]
        pct = count / stats['total_papers'] * 100
        report.append(f"| {year} | {count} | {pct:.1f}% |")

    report.append("\n### 1.2 期刊分布\n")
    report.append("| 期刊名称 | 文献数量 | 占比 |")
    report.append("|----------|----------|------|")
    for journal, count in stats['journal_distribution'].most_common(10):
        pct = count / stats['total_papers'] * 100
        report.append(f"| {journal} | {count} | {pct:.1f}% |")

    # 2. 研究对象与样品类型
    report.append("\n## 2. 研究对象与样品类型\n")
    report.append("| 样品类型 | 文献数量 | 占比 |")
    report.append("|----------|----------|------|")
    for sample, count in sorted(stats['sample_type'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            sample_name = {
                'surface water': '地表水',
                'groundwater': '地下水',
                'drinking water': '饮用水',
                'wastewater': '污水/废水',
                'soil': '土壤/沉积物',
                'biota': '生物样品',
            }.get(sample, sample)
            report.append(f"| {sample_name} | {count} | {pct:.1f}% |")

    # 3. 分析平台与仪器类型
    report.append("\n## 3. 分析平台与仪器类型\n")

    report.append("### 3.1 高分辨率质谱平台\n")
    report.append("| 质谱平台 | 文献数量 | 占比 |")
    report.append("|----------|----------|------|")
    for platform, count in sorted(stats['platform'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            report.append(f"| {platform} | {count} | {pct:.1f}% |")

    report.append("\n### 3.2 色谱联用技术\n")
    report.append("| 联用技术 | 文献数量 | 占比 |")
    report.append("|----------|----------|------|")
    for chrom, count in sorted(stats['chromatography'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            report.append(f"| {chrom} | {count} | {pct:.1f}% |")

    # 4. 筛查策略与分析流程
    report.append("\n## 4. 筛查策略与分析流程\n")
    report.append("| 筛查策略 | 文献数量 | 占比 |")
    report.append("|----------|----------|------|")
    for strategy, count in sorted(stats['screening_strategy'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            strategy_name = {
                'target analysis': '靶向分析',
                'suspect screening': '可疑筛查',
                'non-target screening': '非靶向筛查',
                'combined strategy': '联合策略',
            }.get(strategy, strategy)
            report.append(f"| {strategy_name} | {count} | {pct:.1f}% |")

    # 5. 数据处理与化合物鉴定方法
    report.append("\n## 5. 数据处理与化合物鉴定方法\n")

    report.append("### 5.1 数据处理软件\n")
    report.append("| 软件名称 | 文献数量 | 占比 |")
    report.append("|----------|----------|------|")
    for software, count in sorted(stats['data_software'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            report.append(f"| {software} | {count} | {pct:.1f}% |")

    report.append("\n### 5.2 数据库\n")
    report.append("| 数据库名称 | 文献数量 | 占比 |")
    report.append("|------------|----------|------|")
    for db, count in sorted(stats['database'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            report.append(f"| {db} | {count} | {pct:.1f}% |")

    report.append("\n### 5.3 Schymanski鉴定等级应用情况\n")
    report.append("| 鉴定等级 | 文献数量 | 占比 |")
    report.append("|----------|----------|------|")
    for level, count in sorted(stats['identification_level'].items()):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            level_desc = {
                'Schymanski Level 1': 'Level 1 - 确认结构（参考标准品）',
                'Schymanski Level 2': 'Level 2 - 可能结构（谱库匹配）',
                'Schymanski Level 3': 'Level 3 - 候选结构（化合物类别）',
                'Schymanski Level 4': 'Level 4 - 明确分子式',
                'Schymanski Level 5': 'Level 5 - 精确质量数',
            }.get(level, level)
            report.append(f"| {level_desc} | {count} | {pct:.1f}% |")

    # 6. 污染物类型及应用成果
    report.append("\n## 6. 污染物类型及应用成果\n")
    report.append("| 污染物类型 | 文献数量 | 占比 |")
    report.append("|------------|----------|------|")
    for pollutant, count in sorted(stats['pollutant_type'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            pollutant_name = {
                'PFAS': '全氟和多氟烷基物质（PFAS）',
                'pharmaceuticals': '药物和个人护理品（PPCPs）',
                'pesticides': '农药',
                'endocrine disruptors': '内分泌干扰物',
                'flame retardants': '阻燃剂',
                'industrial chemicals': '工业添加剂',
            }.get(pollutant, pollutant)
            report.append(f"| {pollutant_name} | {count} | {pct:.1f}% |")

    # 7. 技术优势、存在问题与发展趋势
    report.append("\n## 7. 技术优势、存在问题与发展趋势\n")

    report.append("### 7.1 技术优势\n")
    report.append("根据文献分析，HRMS非靶向筛查技术的主要优势包括：\n")
    report.append("- **高分辨率**：能够精确测定化合物的精确质量数，实现未知化合物的初步鉴定")
    report.append("- **高灵敏度**：可检测痕量水平的污染物，满足环境监测需求")
    report.append("- **宽筛查范围**：单次分析可同时检测数百至数千种化合物，包括未知化合物")
    report.append("- **非靶向能力**：无需预先知道目标化合物，可发现新的污染物")
    report.append("- **多策略整合**：可结合靶向、可疑和非靶向筛查策略，提高检出率\n")

    report.append("### 7.2 存在问题\n")
    report.append("当前HRMS非靶向筛查技术面临的主要挑战：\n")
    report.append("- **数据处理复杂**：海量数据的处理和解析需要专业软件和算法")
    report.append("- **标准品缺乏**：许多新污染物缺乏标准品，难以进行准确定量和确认")
    report.append("- **数据库不完善**：现有谱图数据库覆盖范围有限，许多化合物无法匹配")
    report.append("- **未知化合物鉴定困难**：从质谱数据推断未知化合物结构仍然具有挑战性")
    report.append("- **标准化不足**：不同实验室之间的方法和数据处理流程缺乏统一标准\n")

    report.append("### 7.3 发展趋势\n")
    report.append("未来HRMS非靶向筛查技术的发展方向：\n")
    report.append("- **人工智能与机器学习**：利用AI/ML技术提高化合物识别和鉴定的准确性和效率")
    report.append("- **自动化数据解析**：开发自动化数据处理流程，减少人工干预")
    report.append("- **大型谱图库构建**：建立更全面的质谱数据库，扩大化合物覆盖范围")
    report.append("- **标准化流程**：制定统一的分析和数据处理标准，提高结果的可比性")
    report.append("- **多组学整合**：将非靶向筛查与其他组学技术（如代谢组学）结合，提供更全面的信息\n")

    # 8. 文献列表
    report.append("\n## 8. 文献列表\n")
    report.append("| 序号 | 文件名 | 年份 | 领域相关度 |")
    report.append("|------|--------|------|------------|")
    for i, paper in enumerate(data, 1):
        meta = paper.get('metadata', {})
        filename = meta.get('filename', 'Unknown')
        year = meta.get('year', 'Unknown')
        relevance = meta.get('domain_relevance', 0)
        report.append(f"| {i} | {filename[:50]}... | {year} | {relevance:.1%} |")

    # 写入文件
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))

    return output_path


def generate_word_report(data, stats, output_path):
    """生成Word报告"""
    if not HAS_DOCX:
        return None

    doc = Document()

    # 设置标题样式
    title = doc.add_heading('高分辨率质谱非靶向筛查（HRMS-NTS）文献统计报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加基本信息
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'文献总数: {stats["total_papers"]} 篇')

    # 1. 文献基本信息
    doc.add_heading('1. 文献基本信息', level=1)

    doc.add_heading('1.1 发表年份分布', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '年份'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    for year in sorted(stats['year_distribution'].keys()):
        count = stats['year_distribution'][year]
        pct = count / stats['total_papers'] * 100
        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = str(count)
        row_cells[2].text = f'{pct:.1f}%'

    doc.add_heading('1.2 期刊分布', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '期刊名称'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    for journal, count in stats['journal_distribution'].most_common(10):
        pct = count / stats['total_papers'] * 100
        row_cells = table.add_row().cells
        row_cells[0].text = journal
        row_cells[1].text = str(count)
        row_cells[2].text = f'{pct:.1f}%'

    # 2. 研究对象与样品类型
    doc.add_heading('2. 研究对象与样品类型', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '样品类型'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    sample_name_map = {
        'surface water': '地表水',
        'groundwater': '地下水',
        'drinking water': '饮用水',
        'wastewater': '污水/废水',
        'soil': '土壤/沉积物',
        'biota': '生物样品',
    }

    for sample, count in sorted(stats['sample_type'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = sample_name_map.get(sample, sample)
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    # 3. 分析平台与仪器类型
    doc.add_heading('3. 分析平台与仪器类型', level=1)

    doc.add_heading('3.1 高分辨率质谱平台', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '质谱平台'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    for platform, count in sorted(stats['platform'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = platform
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    doc.add_heading('3.2 色谱联用技术', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '联用技术'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    for chrom, count in sorted(stats['chromatography'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = chrom
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    # 4. 筛查策略与分析流程
    doc.add_heading('4. 筛查策略与分析流程', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '筛查策略'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    strategy_name_map = {
        'target analysis': '靶向分析',
        'suspect screening': '可疑筛查',
        'non-target screening': '非靶向筛查',
        'combined strategy': '联合策略',
    }

    for strategy, count in sorted(stats['screening_strategy'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = strategy_name_map.get(strategy, strategy)
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    # 5. 数据处理与化合物鉴定方法
    doc.add_heading('5. 数据处理与化合物鉴定方法', level=1)

    doc.add_heading('5.1 数据处理软件', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '软件名称'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    for software, count in sorted(stats['data_software'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = software
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    doc.add_heading('5.2 数据库', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '数据库名称'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    for db, count in sorted(stats['database'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = db
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    doc.add_heading('5.3 Schymanski鉴定等级应用情况', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '鉴定等级'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    level_desc_map = {
        'Schymanski Level 1': 'Level 1 - 确认结构（参考标准品）',
        'Schymanski Level 2': 'Level 2 - 可能结构（谱库匹配）',
        'Schymanski Level 3': 'Level 3 - 候选结构（化合物类别）',
        'Schymanski Level 4': 'Level 4 - 明确分子式',
        'Schymanski Level 5': 'Level 5 - 精确质量数',
    }

    for level, count in sorted(stats['identification_level'].items()):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = level_desc_map.get(level, level)
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    # 6. 污染物类型及应用成果
    doc.add_heading('6. 污染物类型及应用成果', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '污染物类型'
    hdr_cells[1].text = '文献数量'
    hdr_cells[2].text = '占比'

    pollutant_name_map = {
        'PFAS': '全氟和多氟烷基物质（PFAS）',
        'pharmaceuticals': '药物和个人护理品（PPCPs）',
        'pesticides': '农药',
        'endocrine disruptors': '内分泌干扰物',
        'flame retardants': '阻燃剂',
        'industrial chemicals': '工业添加剂',
    }

    for pollutant, count in sorted(stats['pollutant_type'].items(), key=lambda x: -x[1]):
        if count > 0:
            pct = count / stats['total_papers'] * 100
            row_cells = table.add_row().cells
            row_cells[0].text = pollutant_name_map.get(pollutant, pollutant)
            row_cells[1].text = str(count)
            row_cells[2].text = f'{pct:.1f}%'

    # 7. 技术优势、存在问题与发展趋势
    doc.add_heading('7. 技术优势、存在问题与发展趋势', level=1)

    doc.add_heading('7.1 技术优势', level=2)
    doc.add_paragraph('根据文献分析，HRMS非靶向筛查技术的主要优势包括：')
    advantages = [
        '高分辨率：能够精确测定化合物的精确质量数，实现未知化合物的初步鉴定',
        '高灵敏度：可检测痕量水平的污染物，满足环境监测需求',
        '宽筛查范围：单次分析可同时检测数百至数千种化合物，包括未知化合物',
        '非靶向能力：无需预先知道目标化合物，可发现新的污染物',
        '多策略整合：可结合靶向、可疑和非靶向筛查策略，提高检出率',
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')

    doc.add_heading('7.2 存在问题', level=2)
    doc.add_paragraph('当前HRMS非靶向筛查技术面临的主要挑战：')
    problems = [
        '数据处理复杂：海量数据的处理和解析需要专业软件和算法',
        '标准品缺乏：许多新污染物缺乏标准品，难以进行准确定量和确认',
        '数据库不完善：现有谱图数据库覆盖范围有限，许多化合物无法匹配',
        '未知化合物鉴定困难：从质谱数据推断未知化合物结构仍然具有挑战性',
        '标准化不足：不同实验室之间的方法和数据处理流程缺乏统一标准',
    ]
    for prob in problems:
        doc.add_paragraph(prob, style='List Bullet')

    doc.add_heading('7.3 发展趋势', level=2)
    doc.add_paragraph('未来HRMS非靶向筛查技术的发展方向：')
    trends = [
        '人工智能与机器学习：利用AI/ML技术提高化合物识别和鉴定的准确性和效率',
        '自动化数据解析：开发自动化数据处理流程，减少人工干预',
        '大型谱图库构建：建立更全面的质谱数据库，扩大化合物覆盖范围',
        '标准化流程：制定统一的分析和数据处理标准，提高结果的可比性',
        '多组学整合：将非靶向筛查与其他组学技术（如代谢组学）结合，提供更全面的信息',
    ]
    for trend in trends:
        doc.add_paragraph(trend, style='List Bullet')

    # 8. 文献列表
    doc.add_heading('8. 文献列表', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '文件名'
    hdr_cells[2].text = '年份'
    hdr_cells[3].text = '领域相关度'

    for i, paper in enumerate(data, 1):
        meta = paper.get('metadata', {})
        filename = meta.get('filename', 'Unknown')
        year = meta.get('year', 'Unknown')
        relevance = meta.get('domain_relevance', 0)
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = filename[:50]
        row_cells[2].text = str(year)
        row_cells[3].text = f'{relevance:.1%}'

    # 保存文档
    doc.save(output_path)
    return output_path


def main():
    """主函数"""
    # 配置路径
    json_path = r'd:\VScode\firstcc\GHGs-WWTPs\output\literature_learning\all_papers_analysis.json'
    output_dir = r'd:\VScode\firstcc\GHGs-WWTPs\output\literature_learning'

    # 加载数据
    print("加载分析数据...")
    data = load_analysis_data(json_path)
    print(f"共加载 {len(data)} 篇文献")

    # 提取统计数据
    print("提取统计数据...")
    stats = extract_statistics(data)

    # 生成Markdown报告
    md_path = os.path.join(output_dir, 'HRMS-NTS文献统计报告.md')
    print(f"生成Markdown报告: {md_path}")
    generate_markdown_report(data, stats, md_path)

    # 生成Word报告
    if HAS_DOCX:
        docx_path = os.path.join(output_dir, 'HRMS-NTS文献统计报告.docx')
        print(f"生成Word报告: {docx_path}")
        generate_word_report(data, stats, docx_path)

    print("\n报告生成完成!")
    print(f"Markdown报告: {md_path}")
    if HAS_DOCX:
        print(f"Word报告: {docx_path}")


if __name__ == '__main__':
    main()
