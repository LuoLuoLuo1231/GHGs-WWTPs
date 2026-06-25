"""
Generate formatted Word document from paper sections
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

sys.stdout.reconfigure(encoding='utf-8')

INPUT_DIR = r'D:\VScode\firstcc\GHGs-WWTPs\output\paper_output_v2'
OUTPUT_PATH = os.path.join(INPUT_DIR, 'WWTP_GHG_Methodology_Comparison.docx')

def set_style(doc):
    """Set up academic paper styles"""
    # Title style
    style = doc.styles['Title']
    font = style.font
    font.size = Pt(16)
    font.bold = True
    font.name = 'Times New Roman'

    # Heading 1
    style = doc.styles['Heading 1']
    font = style.font
    font.size = Pt(14)
    font.bold = True
    font.name = 'Times New Roman'

    # Heading 2
    style = doc.styles['Heading 2']
    font = style.font
    font.size = Pt(12)
    font.bold = True
    font.name = 'Times New Roman'

    # Normal
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    font.name = 'Times New Roman'

def add_title_page(doc):
    """Add title page"""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Comparison of Greenhouse Gas Accounting Methodologies\nfor Wastewater Treatment Plants:\nA Meta-Analytical Framework')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('Author Name')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('June 2026')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

def add_abstract(doc, text):
    """Add abstract section"""
    # Abstract title
    h = doc.add_heading('Abstract', level=1)

    # Parse and add abstract text
    lines = text.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('**Keywords'):
            # Keywords line
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', '').replace('*', ''))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.italic = True
        elif line.startswith('**'):
            # Bold section header (Background, Objectives, etc.)
            p = doc.add_paragraph()
            # Extract bold part and normal part
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.font.bold = True
                else:
                    run = p.add_run(part)
        else:
            p = doc.add_paragraph(line)

    doc.add_page_break()

def add_section_from_file(doc, filepath, level=1):
    """Add a section from markdown file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()

    lines = text.split('\n')
    current_para = None

    for line in lines:
        line = line.rstrip()

        # Skip frontmatter
        if line.startswith('---'):
            continue

        # Headings
        if line.startswith('# '):
            heading_text = line[2:].strip()
            doc.add_heading(heading_text, level=level)
            current_para = None
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=level + 1)
            current_para = None
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=level + 2)
            current_para = None
        elif line.startswith('**') and line.endswith('**'):
            # Bold paragraph
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.font.bold = True
            current_para = None
        elif line.startswith('- '):
            # Bullet point
            item_text = line[2:].strip()
            p = doc.add_paragraph(item_text, style='List Bullet')
            current_para = None
        elif line.startswith('[') and '](' in line:
            # Link - skip
            continue
        elif line.strip() == '':
            # Empty line
            current_para = None
        elif line.startswith('|'):
            # Table row - skip for now
            continue
        elif line.startswith('```'):
            # Code block
            continue
        else:
            # Regular paragraph
            if current_para is None:
                current_para = doc.add_paragraph(line)
            else:
                # Continue current paragraph
                current_para.add_run(' ' + line)

def add_references(doc, refs_text):
    """Add references section"""
    doc.add_heading('References', level=1)

    lines = refs_text.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line or line.startswith('#'):
            continue
        if line.startswith('['):
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

def main():
    print('Generating Word document...')

    doc = Document()
    set_style(doc)

    # Title page
    add_title_page(doc)

    # Abstract
    abstract_path = os.path.join(INPUT_DIR, '05_abstract.md')
    if os.path.exists(abstract_path):
        with open(abstract_path, 'r', encoding='utf-8') as f:
            abstract_text = f.read()
        # Remove the heading
        abstract_text = re.sub(r'^# Abstract\s*\n', '', abstract_text)
        add_abstract(doc, abstract_text)
        print('  Added: Abstract')

    # Introduction
    intro_path = os.path.join(INPUT_DIR, '01_introduction.md')
    if os.path.exists(intro_path):
        add_section_from_file(doc, intro_path)
        print('  Added: Introduction')

    # Methods
    methods_path = os.path.join(INPUT_DIR, '02_methods.md')
    if os.path.exists(methods_path):
        add_section_from_file(doc, methods_path)
        print('  Added: Methods')

    # Results
    results_path = os.path.join(INPUT_DIR, '03_results.md')
    if os.path.exists(results_path):
        add_section_from_file(doc, results_path)
        print('  Added: Results')

    # Discussion
    disc_path = os.path.join(INPUT_DIR, '04_discussion.md')
    if os.path.exists(disc_path):
        add_section_from_file(doc, disc_path)
        print('  Added: Discussion')

    # Conclusion
    conc_path = os.path.join(INPUT_DIR, '06_conclusion.md')
    if os.path.exists(conc_path):
        add_section_from_file(doc, conc_path)
        print('  Added: Conclusion')

    # Save
    doc.save(OUTPUT_PATH)
    print(f'\nWord document saved: {OUTPUT_PATH}')
    print(f'File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB')

if __name__ == '__main__':
    main()
